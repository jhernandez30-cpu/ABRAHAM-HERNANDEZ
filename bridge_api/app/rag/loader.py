from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.config import settings
from app.models.schemas import SourceFile


LOGGER = logging.getLogger(__name__)


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    metadata: dict[str, Any]


class DocumentLoader:
    ignored_dirs = {
        ".git",
        ".venv",
        ".obsidian",
        "__pycache__",
        "node_modules",
        ".mypy_cache",
        ".pytest_cache",
    }

    def __init__(self, knowledge_dir: Path | None = None) -> None:
        self.knowledge_dir = knowledge_dir or settings.knowledge_dir

    def list_files(self) -> list[SourceFile]:
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        files: list[SourceFile] = []
        for path in self._iter_supported_files():
            stat = path.stat()
            files.append(
                SourceFile(
                    name=path.name,
                    relative_path=self._relative_path(path),
                    path=str(path),
                    extension=path.suffix.lower(),
                    size=stat.st_size,
                    modified_at=datetime.fromtimestamp(stat.st_mtime, timezone.utc),
                )
            )
        return files

    def load_documents(self, limit: int | None = None) -> tuple[int, list[LoadedDocument]]:
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        documents: list[LoadedDocument] = []
        files_seen = 0
        for path in self._iter_supported_files():
            if limit is not None and files_seen >= limit:
                break
            files_seen += 1
            try:
                documents.extend(self._load_file(path))
            except Exception as exc:
                LOGGER.warning("Could not load document %s: %s", path, exc)
        LOGGER.info("Loaded %s document units from %s files", len(documents), files_seen)
        return files_seen, documents

    def _iter_supported_files(self):
        if not self.knowledge_dir.exists():
            return
        for path in sorted(self.knowledge_dir.rglob("*"), key=self._sort_key):
            if not path.is_file():
                continue
            if any(part in self.ignored_dirs for part in path.parts):
                continue
            if path.suffix.lower() not in settings.allowed_extensions:
                continue
            yield path

    def _sort_key(self, path: Path) -> tuple[int, str]:
        relative = self._relative_path(path)
        first_part = relative.split("/", 1)[0]
        archive_priority = 1 if first_part.startswith("_") else 0
        return archive_priority, relative.lower()

    def _load_file(self, path: Path) -> list[LoadedDocument]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix == ".docx":
            return self._load_docx(path)
        return [LoadedDocument(text=self._load_text(path), metadata=self._base_metadata(path))]

    def _load_pdf(self, path: Path) -> list[LoadedDocument]:
        import fitz

        documents: list[LoadedDocument] = []
        with fitz.open(path) as pdf:
            for page_index, page in enumerate(pdf, start=1):
                text = self._normalize_text(page.get_text("text") or "")
                if not text:
                    continue
                metadata = self._base_metadata(path)
                metadata["page"] = page_index
                documents.append(LoadedDocument(text=text, metadata=metadata))
        return documents

    def _load_docx(self, path: Path) -> list[LoadedDocument]:
        from docx import Document

        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        text = self._normalize_text("\n".join(paragraphs))
        if not text:
            return []
        return [LoadedDocument(text=text, metadata=self._base_metadata(path))]

    def _load_text(self, path: Path) -> str:
        raw = path.read_text(encoding="utf-8", errors="replace")
        if path.suffix.lower() == ".json":
            raw = self._pretty_json(raw)
        return self._normalize_text(raw)

    def _pretty_json(self, raw: str) -> str:
        try:
            return json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return raw

    def _base_metadata(self, path: Path) -> dict[str, Any]:
        stat = path.stat()
        return {
            "source": self._relative_path(path),
            "source_path": str(path),
            "file_name": path.name,
            "title": path.stem,
            "file_type": path.suffix.lower(),
            "modified_at": int(stat.st_mtime),
            "size": int(stat.st_size),
        }

    def _relative_path(self, path: Path) -> str:
        try:
            return path.relative_to(self.knowledge_dir).as_posix()
        except ValueError:
            return path.name

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\x00", " ")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        return text.strip()
